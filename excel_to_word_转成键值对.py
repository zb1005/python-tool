import pandas as pd
import json
from docx import Document
import os

def excel_to_word(input_excel, output_word):
    # 读取Excel文件
    df = pd.read_excel(input_excel)

    # 创建Word文档
    doc = Document()
    # 添加标题
    doc.add_heading(input_excel.split("\\")[-1].split(".")[0] + "对照关系", level=1)

    # 添加复合键值对内容
    doc.add_paragraph('以下是' + input_excel.split("\\")[-1].split(".")[0] + '的对照关系:')
    doc.add_paragraph('')  # 添加空行

    #获取表格的所有列名
    col_names = df.columns.tolist()

    # 遍历数据行并添加键值对
    for _, row in df.iterrows():
        # 创建复合键值对文本
        # 创建单行JSON格式的键值对
        data = {}
        for col_name in col_names:
            data[col_name] = str(row[col_name])
        key_value_text = json.dumps(data, ensure_ascii=False, separators=(',', ':'))
        
        # 添加键值对到文档
        doc.add_paragraph(key_value_text)
        doc.add_paragraph('---')  # 添加分隔线

    # 添加说明文本
    # doc.add_paragraph('注：以上内容已按要求格式转换为复合键值对形式，可直接用于AI知识库构建。')

    # 保存文档
    doc.save(output_word)
    print(f'Word文档已生成: {output_word}')

if __name__ == "__main__":
    # input_files = [
    #     r"C:\Users\zhangbon\Desktop\2025_06_12_AI知识库\AI知识库\数据消费\元数据字典值集地图_值集.xlsx",
    #     r"C:\Users\zhangbon\Desktop\2025_06_12_AI知识库\AI知识库\数据消费\元数据字典值集地图_元数据字典.xlsx",
    #     r"C:\Users\zhangbon\Desktop\2025_06_12_AI知识库\AI知识库\数据消费\元数据字典值集地图_元数据字典值集.xlsx"]
    #input_files为获取一个文件夹下面的所有的文件
    # input_files_name=os.listdir(r"C:\Users\zhangbon\Desktop\2025_06_12_AI知识库\AI知识库\数据消费副本")
    input_files_name=os.listdir(r"C:\Users\zhangbon\Desktop\洗")
    #input_files为筛选出所有的xlsx文件
    input_files=[os.path.join(r"C:\Users\zhangbon\Desktop\洗",input_file) for input_file in input_files_name if input_file.endswith(".xlsx") and not input_file.startswith('~$')]
    output_files = [input_file.replace(".xlsx","对照关系.docx") for input_file in input_files]
    for input_file,output_file in zip(input_files,output_files):
        print(f"=======当前处理{input_file}=====")
        excel_to_word(input_file, output_file)
