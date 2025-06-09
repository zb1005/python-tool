import pandas as pd
from docx import Document

def excel_to_word(input_excel, output_word):
    # 读取Excel文件
    df = pd.read_excel(input_excel)
    
    # 检查列是否存在
    required_columns = ['问题描述', '解决方案（要足够详细）']
    if not all(col in df.columns for col in required_columns):
        print(f"错误：Excel文件中必须包含{required_columns}两列")
        return
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('业务问题与解决方案汇总', level=1)
    
    # 遍历Excel中的每一行数据
    for _, row in df.iterrows():
        # 添加问题描述
        doc.add_paragraph("问题描述：", style='Heading 3')
        doc.add_paragraph(str(row['问题描述']) if pd.notna(row['问题描述']) else "无")
        
        # 添加解决方案
        doc.add_paragraph("解决方案：", style='Heading 3')
        solution = row['解决方案（要足够详细）']
        doc.add_paragraph(str(solution) if pd.notna(solution) else "无")
        
        # 添加分隔线
        doc.add_paragraph("_________________________")
    
    # 保存Word文档
    doc.save(output_word)
    print(f"Word文档已成功生成：{output_word}")

if __name__ == "__main__":
    input_file = r"C:\Users\zhangbon\Desktop\业务问题&需求汇总表.xlsx"
    output_file = r"C:\Users\zhangbon\Desktop\业务问题&需求汇总表.docx"
    excel_to_word(input_file, output_file)