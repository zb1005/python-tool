import os
import argparse
from pdf2docx import Converter
import logging
from docx import Document
from docx.shared import Pt
import glob

def setup_logging():
    """设置日志配置，同时输出到文件和控制台"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('pdf_to_word_conversion.log'),
            logging.StreamHandler()
        ]
    )

def pdf_to_word(pdf_path, output_dir=None):
    """
    将单个PDF文件转换为Word文档
    
    Args:
        pdf_path (str): PDF文件的路径
        output_dir (str, optional): 输出Word文档的目录. 
                                 如果为None，则与PDF文件在同一目录
        
    Returns:
        bool: 转换成功返回True，否则返回False
    """
    try:
        # 检查PDF文件是否存在
        if not os.path.exists(pdf_path):
            logging.error(f"文件不存在: {pdf_path}")
            return False
        
        # 获取PDF文件名（不含扩展名）
        pdf_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        
        # 确定输出目录
        if output_dir is None:
            output_dir = os.path.dirname(pdf_path)
        
        # 创建输出目录（如果不存在）
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成Word文件路径
        word_path = os.path.join(output_dir, f"{pdf_filename}.docx")
        
        # 检查Word文件是否已存在
        if os.path.exists(word_path):
            logging.warning(f"Word文件已存在，将被覆盖: {word_path}")
        
        # 转换PDF到Word
        logging.info(f"开始转换: {pdf_path} -> {word_path}")
        cv = Converter(pdf_path)
        cv.convert(word_path, start=0, end=None)  # 转换所有页面
        cv.close()
        
        logging.info(f"转换成功: {word_path}")
        return True
    
    except Exception as e:
        logging.error(f"转换失败: {pdf_path}，错误信息: {str(e)}")
        return False

def batch_convert_pdfs(input_dir, output_dir=None, recursive=False):
    """
    批量转换目录中的PDF文件为Word文档
    
    Args:
        input_dir (str): 包含PDF文件的目录
        output_dir (str, optional): 输出Word文档的目录.
                                 如果为None，则与PDF文件在同一目录
        recursive (bool, optional): 是否递归处理子目录. 默认是False
        
    Returns:
        tuple: (成功转换的数量, 总PDF文件数量)
    """
    # 检查输入目录是否存在
    if not os.path.isdir(input_dir):
        logging.error(f"输入目录不存在: {input_dir}")
        return (0, 0)
    
    total = 0
    success = 0
    
    # 遍历目录中的文件
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            # 检查文件是否为PDF
            if file.lower().endswith('.pdf'):
                total += 1
                pdf_path = os.path.join(root, file)
                
                # 确定当前文件的输出目录
                if output_dir is not None:
                    # 如果递归处理，保持目录结构
                    if recursive:
                        relative_path = os.path.relpath(root, input_dir)
                        current_output_dir = os.path.join(output_dir, relative_path)
                    else:
                        current_output_dir = output_dir
                else:
                    current_output_dir = None
                
                # 转换PDF到Word
                if pdf_to_word(pdf_path, current_output_dir):
                    success += 1
        
        # 如果不递归处理子目录，只处理当前目录后就退出
        if not recursive:
            break
    
    logging.info(f"批量转换完成。成功: {success}/{total}")
    return (success, total)

def read_order_list(order_file_path):
    """
    读取顺序列表文件，返回PDF文件路径列表
    
    Args:
        order_file_path (str): 顺序列表文件路径
        
    Returns:
        list: PDF文件路径列表，按照顺序排列
    """
    if not os.path.exists(order_file_path):
        logging.error(f"顺序列表文件不存在: {order_file_path}")
        return []
    
    pdf_paths = []
    try:
        with open(order_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):  # 跳过空行和注释
                    pdf_paths.append(line)
        logging.info(f"成功读取顺序列表，共 {len(pdf_paths)} 个PDF文件")
        return pdf_paths
    except Exception as e:
        logging.error(f"读取顺序列表失败: {str(e)}")
        return []

def merge_word_documents(word_files, output_path):
    """
    合并多个Word文档到一个文档中，每个文档前添加文件名作为标题
    
    Args:
        word_files (list): Word文件路径列表
        output_path (str): 合并后的输出文件路径
        
    Returns:
        bool: 合并成功返回True，否则返回False
    """
    try:
        merged_doc = Document()
        
        for i, word_file in enumerate(word_files):
            if os.path.exists(word_file):
                doc = Document(word_file)
                
                # 如果不是第一个文档，添加分页符
                if i > 0:
                    merged_doc.add_page_break()
                
                # 添加文件名作为标题
                filename = os.path.splitext(os.path.basename(word_file))[0]
                title_paragraph = merged_doc.add_paragraph()
                title_run = title_paragraph.add_run(f"文档标题: {filename}")
                title_run.bold = True
                title_run.font.size = Pt(14)
                merged_doc.add_paragraph()  # 空行
                
                # 复制所有段落
                for paragraph in doc.paragraphs:
                    new_paragraph = merged_doc.add_paragraph()
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                
                # 复制所有表格
                for table in doc.tables:
                    new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.cell(i, j).text = cell.text
                
                logging.info(f"已合并: {word_file}")
            else:
                logging.warning(f"Word文件不存在，跳过: {word_file}")
        
        # 保存合并后的文档
        merged_doc.save(output_path)
        logging.info(f"合并完成: {output_path}")
        return True
        
    except Exception as e:
        logging.error(f"合并Word文档失败: {str(e)}")
        return False

def convert_with_order(order_file_path, output_dir):
    """
    按照顺序列表转换PDF文件并合并为一个Word文档
    
    Args:
        order_file_path (str): 顺序列表文件路径
        output_dir (str): 输出目录
        
    Returns:
        tuple: (成功转换的数量, 总PDF文件数量, 合并后的Word文件路径)
    """
    # 读取顺序列表
    pdf_paths = read_order_list(order_file_path)
    if not pdf_paths:
        return (0, 0, None)
    
    total = len(pdf_paths)
    success = 0
    converted_files = []
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 按照顺序转换每个PDF文件
    for i, pdf_path in enumerate(pdf_paths, 1):
        logging.info(f"正在处理第 {i}/{total} 个文件: {pdf_path}")
        
        # 转换PDF到Word
        if pdf_to_word(pdf_path, output_dir):
            success += 1
            # 获取转换后的Word文件路径
            pdf_filename = os.path.splitext(os.path.basename(pdf_path))[0]
            word_path = os.path.join(output_dir, f"{pdf_filename}.docx")
            converted_files.append(word_path)
    
    # 合并所有转换后的Word文档
    merged_output_path = os.path.join(output_dir, "merged_document.docx")
    if converted_files:
        merge_success = merge_word_documents(converted_files, merged_output_path)
        if merge_success:
            logging.info(f"所有文件已成功合并到: {merged_output_path}")
        else:
            logging.error("合并Word文档失败")
    
    logging.info(f"顺序转换完成。成功: {success}/{total}")
    return (success, total, merged_output_path)

def main():
    # 设置命令行参数解析
    parser = argparse.ArgumentParser(description='按照顺序列表将PDF文件转换为Word并合并')
    parser.add_argument('--order-file', '-o', required=True, help='顺序列表文件路径')
    parser.add_argument('--output-dir', '-d', required=True, help='输出目录')
    parser.add_argument('--create-sample', action='store_true', help='创建示例顺序列表文件')
    
    args = parser.parse_args()
    
    # 设置日志
    setup_logging()
    
    # 如果用户要求创建示例顺序列表文件
    if args.create_sample:
        sample_content = """# PDF文件顺序列表
# 每行一个PDF文件路径，按顺序排列
# 注释以#开头

D:\\path\\to\\file1.pdf
D:\\path\\to\\file2.pdf
D:\\path\\to\\file3.pdf
"""
        sample_path = "sample_order_list.txt"
        with open(sample_path, 'w', encoding='utf-8') as f:
            f.write(sample_content)
        print(f"已创建示例顺序列表文件: {sample_path}")
        return
    
    # 执行顺序转换和合并
    success, total, merged_file = convert_with_order(args.order_file, args.output_dir)
    
    # 输出结果摘要
    print(f"转换完成: {success} 个成功，{total - success} 个失败，共 {total} 个PDF文件")
    if merged_file:
        print(f"合并后的Word文档: {merged_file}")

if __name__ == "__main__":
    main()
    